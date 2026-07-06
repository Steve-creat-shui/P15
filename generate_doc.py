#!/usr/bin/env python3
"""Generate a one-page Word document introducing the JEVS system."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page margins (narrow to fit more content) ---
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# --- Default style ---
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# --- Helper functions ---
def add_heading_styled(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 0:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif level == 1:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif level == 2:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    heading.paragraph_format.space_before = Pt(6) if level > 0 else Pt(0)
    heading.paragraph_format.space_after = Pt(3)
    return heading

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 + level * 0.6)
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(9.5)
    return p

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
title = doc.add_heading('JEVS 司法证据可视化系统', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(18)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(8)
run = sub.add_run('Judicial Evidence Visualization System')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

# Divider
div = doc.add_paragraph()
div.paragraph_format.space_after = Pt(6)
run = div.add_run('━' * 50)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# --- Section 1: Overview ---
add_heading_styled('一、系统概述', level=1)
add_para(
    'JEVS 是一套面向法学教育的智能教学辅助平台。系统以法院裁判文书、起诉书、案卷材料等非结构化法律文本为输入，'
    '通过 AI 大模型自动提取证据要素，并结合三维空间映射引擎与多模态图像生成技术，将枯燥的文字案件转化为包含'
    '三维犯罪现场重建图、证据特写照片、法医伤情示意图及结构化案件分析报告在内的可视化教学资源，'
    '帮助法学院学生直观理解案件的空间关系、物证分布与证据链逻辑。'
)

# --- Section 2: Problems Solved ---
add_heading_styled('二、解决的核心问题', level=1)

problems = [
    ('法学教育可视化缺失：', '传统教学中，学生只能通过纯文本裁判文书学习案件，缺乏对现场空间关系和物证分布的直观感知。JEVS 将文字转化为图像，让案件变得"看得见"。'),
    ('AI 幻觉防控：', '通用大模型在法律场景中容易虚构证据细节。JEVS 设计了严格的证据过滤器（硬编码规则约束可提取/不可提取内容）和确定性空间映射引擎（无 LLM 参与的坐标映射），从架构层面杜绝 AI 编造证据。'),
    ('教师审核工作流：', 'AI 提取的证据需经教师（超级用户）逐条审核、修改、排除后方可进入场景生成环节，确保教学内容的准确性和权威性。'),
    ('文书证据保真渲染：', '对于合同、借条、聊天记录等文书类证据，系统使用纯 Pillow 代码渲染图像（不使用 AI 生成文字），保证文字内容 100% 准确，杜绝任何虚构。'),
]
for title_text, desc in problems:
    add_bullet(f'{title_text}{desc}')

# --- Section 3: Features ---
add_heading_styled('三、核心功能', level=1)

features = [
    '案件管理：支持粘贴文本或上传 TXT/PDF/DOCX 文件创建案件，自动识别文件编码（UTF-8/GBK），状态流跟踪（待处理→已提取→已审核→已生成）。',
    'AI 证据智能提取：基于 DeepSeek/OpenAI 等大模型 + Instructor 结构化输出框架，自动从裁判文书中分离可可视化证据、待确认证据和不可可视化内容，支持物证状态（血迹形态、指纹、损坏程度等）的结构化描述。',
    '证据审核与管理：教师可逐条审核、编辑、排除证据，将证据分配到不同场景，确保进入图像生成环节的数据完全可控。',
    '三维场景重建：确定性空间映射引擎内置 45+ 位置坐标（如"床下左侧""桌面右上"）到 3×3 空间网格的映射规则，支持 17 种房间类型自动识别，将平面证据列表转化为有空间关系的三维场景描述。',
    '多策略图像生成：支持 DALL·E / FLUX / 混元 / Agnes / ZenMux 五种图像生成后端，可按场景类型独立配置；支持三种证据特写策略——几何裁剪（零成本）、背景修复式生成和纯 AI 锁风格生成。',
    '法医伤情示意图：AI 生成人体部位损伤特写照片，搭配 Pillow 渲染的人体轮廓标注图（带编号和文字说明，自动避让重叠标注）。',
    '文书证据渲染：100% 代码渲染，不依赖 AI——支持方格纸手写体、A4 公文（含圆形红色印章）和微信聊天截图三种模板，中文自动换行。',
    '多场景管理：自动扫描文本关键词建议场景（卧室、客厅、厨房等），支持一个案件创建多个场景独立生成图像，批量生成支持 API 速率限制保护。',
    '案件报告导出：支持 Markdown 和 HTML 格式，自动汇总证据清单、场景状态和原文摘录。教案渲染器可生成结构化教学方案，包含证据链表格、场景图片和课堂讨论问题。',
    '用户权限体系：JWT 认证 + Argon2 密码哈希，超级用户（教师）与普通用户（学生）角色分离，支持邮箱密码找回。',
]
for f in features:
    add_bullet(f)

# --- Section 4: Tech Stack ---
add_heading_styled('四、技术架构', level=1)
add_para(
    '后端：Python 3.10+ / FastAPI + SQLModel (PostgreSQL 18) / Instructor + OpenAI SDK（LLM 结构化输出）/ '
    'Pillow（图像渲染）/ pypdf + zipfile（文档解析）/ JWT + Argon2（认证）/ Docker + Traefik（容器化部署）',
    indent=True
)
add_para(
    '前端：React 19 + TypeScript / Vite / TanStack Router + Query + Table / shadcn/ui (Radix) + Tailwind CSS 4 / '
    'react-hook-form + zod / Framer Motion（动效）',
    indent=True
)
add_para(
    '基础设施：GitHub Actions（CI/CD）/ Sentry（监控）/ SMTP（邮件）/ Nginx + 静态文件服务',
    indent=True
)

# Save
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'JEVS系统介绍.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
